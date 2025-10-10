"""Tests for document processing functionality."""
import pytest
from pathlib import Path
import io
import docx
from app.utils.document_processor import (
    DocumentProcessor,
    FileSizeError,
    FileTypeError,
    DocumentValidationError
)

@pytest.fixture
def processor():
    """Create a DocumentProcessor instance."""
    return DocumentProcessor()

@pytest.fixture
def sample_txt_file(tmp_path):
    """Create a sample text file."""
    file_path = tmp_path / "test.txt"
    file_path.write_text("Sample text content")
    return file_path

@pytest.fixture
def large_txt_file(tmp_path):
    """Create a large text file that exceeds size limit."""
    file_path = tmp_path / "large.txt"
    # Create file larger than MAX_FILE_SIZE
    with open(file_path, 'wb') as f:
        f.write(b'x' * (DocumentProcessor.MAX_FILE_SIZE + 1))
    return file_path

def test_validate_file_with_valid_txt(processor, sample_txt_file):
    """Test file validation with valid text file."""
    mime_type, format_id = processor.validate_file(sample_txt_file)
    assert mime_type == 'text/plain'
    assert format_id == 'txt'

def test_validate_file_size_limit(processor, large_txt_file):
    """Test file size validation."""
    with pytest.raises(FileSizeError):
        processor.validate_file(large_txt_file)

def test_validate_file_not_found(processor):
    """Test validation of non-existent file."""
    with pytest.raises(FileNotFoundError):
        processor.validate_file("nonexistent.txt")

def test_validate_file_unsupported_type(processor, tmp_path):
    """Test validation of unsupported file type."""
    # Create a binary file that's definitely not text/pdf/docx
    file_path = tmp_path / "test.bin"
    with open(file_path, 'wb') as f:
        f.write(b'\x00\x01\x02\x03')  # Binary content
    with pytest.raises(FileTypeError):
        processor.validate_file(file_path)

def test_process_document_txt(processor, sample_txt_file):
    """Test processing of text document."""
    result = processor.process_document(sample_txt_file)
    assert "text" in result
    assert "metadata" in result
    assert result["text"] == "Sample text content"
    assert result["metadata"]["format"] == "txt"
    assert result["metadata"]["words"] == 3

@pytest.fixture
def sample_docx_file(tmp_path):
    """Create a sample DOCX file for testing."""
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph with some content.')
    doc.add_paragraph('This is another paragraph with different content.')
    
    file_path = tmp_path / "test.docx"
    doc.save(str(file_path))
    return file_path

def test_process_document_docx(processor, sample_docx_file):
    """Test processing of DOCX document with enhanced features."""
    result = processor.process_document(sample_docx_file)
    
    # Basic structure checks
    assert "text" in result
    assert "metadata" in result
    assert "paragraphs" in result
    
    # Metadata checks
    metadata = result["metadata"]
    assert "document_info" in metadata
    assert "statistics" in metadata
    assert "formatting" in metadata
    
    # Statistics checks
    stats = metadata["statistics"]
    assert stats["paragraphs"] > 0
    assert stats["words"] > 0
    assert stats["sections"] > 0
    
    # Paragraph structure checks
    assert len(result["paragraphs"]) > 0
    for para in result["paragraphs"]:
        assert "text" in para
        assert "style" in para
        assert "alignment" in para
        assert "font_size" in para
        assert isinstance(para["is_bold"], bool)
        assert isinstance(para["is_italic"], bool)

@pytest.fixture
def sample_pdf_file(tmp_path):
    """Create a sample PDF file for testing."""
    try:
        import reportlab.pdfgen.canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        
        pdf_path = tmp_path / "test.pdf"
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        
        # Create story with different elements
        story = []
        styles = getSampleStyleSheet()
        
        # Add some text
        story.append(Paragraph("Test Document", styles['Heading1']))
        story.append(Paragraph("This is a test paragraph with some content.", styles['Normal']))
        
        # Add a table
        data = [['Header 1', 'Header 2'], ['Cell 1', 'Cell 2']]
        table = Table(data)
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        
        # Build PDF
        doc.build(story)
        return pdf_path
        
    except ImportError:
        pytest.skip("reportlab not installed")

def test_process_document_pdf(processor, sample_pdf_file):
    """Test processing of PDF document with enhanced features."""
    result = processor.process_document(sample_pdf_file)
    
    # Basic structure checks
    assert "text" in result
    assert "metadata" in result
    assert "pages" in result
    
    # Metadata structure checks
    metadata = result["metadata"]
    assert "document_info" in metadata
    assert "statistics" in metadata
    assert "formatting" in metadata
    
    # Statistics checks
    stats = metadata["statistics"]
    assert "pages" in stats
    assert "words" in stats
    assert "characters" in stats
    assert "tables" in stats
    assert "images" in stats
    assert "fonts" in stats
    assert "average_words_per_page" in stats
    
    # Page structure checks
    assert len(result["pages"]) > 0
    for page in result["pages"]:
        assert "page_number" in page
        assert "text" in page
        assert "width" in page
        assert "height" in page
        assert "tables" in page
        assert "images" in page
        assert "words" in page
        assert "characters" in page